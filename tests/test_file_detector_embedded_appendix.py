from docx import Document

from licitasuite.core.file_detector import FileDetector


def test_file_detector_accepts_appendix_table_inside_model(tmp_path):
    model = tmp_path / "ATA DE REGISTRO DE PRECOS - MODELO.docx"
    pdf = tmp_path / "RelatorioItensVencidos.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")

    doc = Document()
    doc.add_paragraph("ATA DE REGISTRO DE PRECOS")
    table = doc.add_table(rows=4, cols=6)
    table.cell(0, 0).text = "Codigo Siplan"
    table.cell(0, 1).text = "Item"
    table.cell(0, 2).text = "Descritivo"
    table.cell(0, 3).text = "Apresentacao"
    table.cell(0, 4).text = "Quantidade"
    table.cell(0, 5).text = "Total"
    table.cell(1, 0).text = "1001"
    table.cell(1, 1).text = "1"
    table.cell(1, 2).text = "Produto um"
    table.cell(1, 3).text = "UN"
    table.cell(1, 4).text = "10"
    table.cell(2, 0).text = "1002"
    table.cell(2, 1).text = "2"
    table.cell(2, 2).text = "Produto dois"
    table.cell(2, 3).text = "CX"
    table.cell(2, 4).text = "20"
    doc.save(model)

    detected = FileDetector(tmp_path).detect()

    assert detected.modelo_ata == model
    assert detected.apendice == model
    assert detected.apendice_embutido is True
    assert detected.vencedores_pdf == pdf
