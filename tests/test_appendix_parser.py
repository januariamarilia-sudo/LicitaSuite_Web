from docx import Document

from licitasuite.parsers.appendix_parser import AppendixParser


def test_appendix_parser_finds_shifted_header_with_merged_cells(tmp_path):
    path = tmp_path / "apendice.docx"
    doc = Document()
    table = doc.add_table(rows=5, cols=7)

    table.cell(0, 0).text = "APENDICE DO PROCESSO"
    table.cell(0, 0).merge(table.cell(0, 6))
    table.cell(1, 0).text = "Codigo Siplan"
    table.cell(1, 1).text = "Item"
    table.cell(1, 2).text = "Descricao\nOficial"
    table.cell(1, 3).text = "Marca"
    table.cell(1, 4).text = "Unidade"
    table.cell(1, 5).text = "Quantidade"
    table.cell(1, 6).text = "Observacao"

    table.cell(2, 0).text = "99123"
    table.cell(2, 1).text = "1"
    table.cell(2, 2).text = "Seringa descartavel"
    table.cell(2, 4).text = "UN"
    table.cell(2, 5).text = "1.250"

    table.cell(3, 0).text = "99124"
    table.cell(3, 1).text = "2"
    table.cell(3, 2).text = "Agulha hipodermica"
    table.cell(3, 4).text = "CX"
    table.cell(3, 5).text = "50"

    doc.save(path)

    parser = AppendixParser()
    itens = parser.parse(path)

    assert [item.numero_item for item in itens] == [1, 2]
    assert itens[0].codigo_siplan == "99123"
    assert itens[0].descricao == "Seringa descartavel"
    assert itens[0].apresentacao == "UN"
    assert itens[0].total == 1250
    assert any("Tabela 1" in line and "selecionada" in line for line in parser.diagnostics)


def test_appendix_parser_falls_back_to_largest_numeric_item_table(tmp_path):
    path = tmp_path / "apendice_fallback.docx"
    doc = Document()
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "Resumo"
    table = doc.add_table(rows=4, cols=4)
    table.cell(0, 0).text = "Linha informativa"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "Codigo 111"
    table.cell(1, 2).text = "Equipo macro gotas"
    table.cell(1, 3).text = "10"
    table.cell(2, 0).text = "2"
    table.cell(2, 1).text = "Codigo 222"
    table.cell(2, 2).text = "Coletor perfurocortante"
    table.cell(2, 3).text = "20"
    doc.save(path)

    parser = AppendixParser()
    itens = parser.parse(path)

    assert [item.numero_item for item in itens] == [1, 2]
    assert itens[0].descricao == "Equipo macro gotas"
    assert any("fallback" in line and "selecionada" in line for line in parser.diagnostics)
