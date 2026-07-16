from pathlib import Path
from docx import Document
from licitasuite.parsers.text_utils import normalize_text

class DetectedFiles:
    def __init__(self, modelo_ata=None, apendice=None, vencedores_pdf=None, banco_fornecedores=None):
        self.modelo_ata = modelo_ata
        self.apendice = apendice
        self.vencedores_pdf = vencedores_pdf
        self.banco_fornecedores = banco_fornecedores

    def missing(self):
        missing = []
        if not self.modelo_ata:
            missing.append("modelo da ata")
        if not self.apendice:
            missing.append("apêndice")
        if not self.vencedores_pdf:
            missing.append("PDF dos vencedores")
        return missing

class FileDetector:
    def __init__(self, folder):
        self.folder = Path(folder)

    def detect(self):
        docs = list(self.folder.rglob("*.docx"))
        pdfs = list(self.folder.rglob("*.pdf"))
        xlsxs = list(self.folder.rglob("*.xlsx"))

        apendice = None
        modelo = None

        for docx in docs:
            name = normalize_text(docx.name)
            if "APENDICE" in name:
                apendice = docx
                continue
            if not modelo and ("ATA" in name or "REGISTRO DE PRECOS" in name):
                modelo = docx
            try:
                doc = Document(docx)
                text = normalize_text("\n".join(p.text for p in doc.paragraphs[:30]))
                table_text = normalize_text(" ".join(cell.text for t in doc.tables[:1] for r in t.rows[:2] for cell in r.cells)) if doc.tables else ""
                if not apendice and ("APENDICE" in text or ("COD" in table_text and "MUNIC" in table_text)):
                    apendice = docx
                elif not modelo and ("ATA DE REGISTRO DE PRECOS" in text or "PROCESSO LICITATORIO" in text):
                    modelo = docx
            except Exception:
                pass

        if not modelo:
            for docx in docs:
                if docx != apendice:
                    modelo = docx
                    break

        vencedores = None
        for pdf in pdfs:
            if "vencedor" in pdf.name.lower():
                vencedores = pdf
                break
        if not vencedores and pdfs:
            vencedores = pdfs[0]

        banco = None
        for xlsx in xlsxs:
            name = xlsx.name.lower()
            if "banco" in name or "fornecedor" in name or "dados" in name:
                banco = xlsx
                break
        if not banco and xlsxs:
            banco = xlsxs[0]

        return DetectedFiles(modelo, apendice, vencedores, banco)
