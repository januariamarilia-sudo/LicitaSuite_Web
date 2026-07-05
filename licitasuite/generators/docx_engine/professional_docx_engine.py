from pathlib import Path
from docx import Document

from licitasuite.generators.docx_engine.document_replacer import DocumentReplacer
from licitasuite.generators.docx_engine.table_locator import TableLocator
from licitasuite.generators.docx_engine.table_writer import TableWriter

class ProfessionalDocxEngine:
    """
    Motor principal para edição do modelo oficial.
    """

    def __init__(self, template_path):
        self.template_path = Path(template_path)

    def open(self):
        return Document(self.template_path)

    def apply_replacements(self, document, replacements):
        return DocumentReplacer().replace(document, replacements)

    def append_items_table(self, document, ata):
        document.add_page_break()
        document.add_heading("QUADRO DE ITENS — LICITASUITE", level=1)

        table = document.add_table(rows=1, cols=9)
        table.style = "Table Grid"

        headers = [
            "CÓDIGO", "ITEM", "DESCRIÇÃO", "MARCA", "FABRICANTE",
            "MODELO", "QTD", "UNITÁRIO", "TOTAL"
        ]

        writer = TableWriter()
        writer.write_row(table.rows[0], headers)

        for item in ata.itens:
            row = table.add_row()
            writer.write_row(row, [
                item.codigo_siplan,
                item.numero_item,
                item.descricao_oficial,
                item.marca,
                item.fabricante,
                item.modelo,
                item.quantidade,
                item.valor_unitario,
                item.valor_total,
            ])

        return document

    def save(self, document, output_path):
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return output_path
