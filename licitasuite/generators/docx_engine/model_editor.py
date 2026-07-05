from docx import Document
from licitasuite.generators.docx_engine.placeholder_replacer import PlaceholderReplacer
from licitasuite.utils.formatters import format_brl, format_quantity
from licitasuite.utils.value_extenso import valor_por_extenso

class ModelEditor:
    def open(self, template_path):
        return Document(template_path)

    def apply_supplier_data(self, document, ata):
        mapping = {
            "{{FORNECEDOR}}": ata.fornecedor_nome,
            "{{RAZAO_SOCIAL}}": ata.fornecedor_nome,
            "{{CNPJ}}": ata.fornecedor_cnpj,
            "{{VALOR_TOTAL}}": format_brl(ata.valor_total, 2),
            "{{VALOR_TOTAL_EXTENSO}}": valor_por_extenso(ata.valor_total),
        }
        return PlaceholderReplacer().replace_document(document, mapping)

    def append_generated_summary(self, document, ata):
        document.add_page_break()
        document.add_heading("ANEXO GERADO PELO LICITASUITE v2.0", level=1)
        document.add_paragraph(f"Fornecedor: {ata.fornecedor_nome}")
        document.add_paragraph(f"CNPJ: {ata.fornecedor_cnpj}")
        document.add_paragraph(f"Valor total: {format_brl(ata.valor_total, 2)} ({valor_por_extenso(ata.valor_total)}).")

        table = document.add_table(rows=1, cols=9)
        table.style = "Table Grid"

        headers = [
            "CÓDIGO SIPLAN", "ITEM", "DESCRIÇÃO OFICIAL", "MARCA",
            "FABRICANTE", "MODELO", "QUANT.", "VALOR UNIT.", "VALOR TOTAL"
        ]

        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header

        for item in ata.itens:
            row = table.add_row().cells
            values = [
                item.codigo_siplan,
                str(item.numero_item),
                item.descricao_oficial,
                item.marca,
                item.fabricante,
                item.modelo,
                format_quantity(item.quantidade),
                format_brl(item.valor_unitario, 4),
                format_brl(item.valor_total, 2),
            ]
            for idx, value in enumerate(values):
                row[idx].text = str(value)

        if ata.inconsistencias:
            document.add_paragraph("")
            document.add_paragraph("INCONSISTÊNCIAS / INFORMAÇÕES NÃO LOCALIZADAS:")
            for inc in ata.inconsistencias:
                document.add_paragraph(f"- {inc}")

        return document
