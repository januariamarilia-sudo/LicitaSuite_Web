from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
from copy import deepcopy
import re
import shutil

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from licitasuite.utils.formatters import safe_filename
from licitasuite.utils.value_extenso import valor_por_extenso

NAO_LOCALIZADA = "[INFORMAÇÃO NÃO LOCALIZADA]"

class ModelBasedAtaGenerator:
    def __init__(self, modelo_path, output_dir="output/atas_geradas"):
        self.modelo_path = Path(modelo_path)
        self.output_dir = Path(output_dir)

    def generate_all(self, atas):
        self.output_dir.mkdir(parents=True, exist_ok=True)
        for old in self.output_dir.glob("*.docx"):
            try:
                old.unlink()
            except Exception:
                pass

        generated = []
        for ata in atas:
            if not ata.itens:
                continue
            generated.append(self.generate_one(ata))

        zip_path = self.output_dir.parent / "atas_geradas.zip"
        if zip_path.exists():
            zip_path.unlink()

        with ZipFile(zip_path, "w", ZIP_DEFLATED) as z:
            for file in generated:
                z.write(file, file.name)

        return generated, zip_path

    def generate_one(self, ata):
        doc = Document(self.modelo_path)

        self.replace_supplier_intro(doc, ata)
        self.replace_registered_prices_table(doc, ata)
        self.replace_total_paragraph(doc, ata)
        self.replace_appendix_table(doc, ata)
        self.replace_signature(doc, ata)

        filename = f"ATA DE REGISTRO DE PREÇOS - {safe_filename(ata.fornecedor_nome)}.docx"
        output = self.output_dir / filename
        doc.save(output)
        return output

    # ------------------------ texto do fornecedor ------------------------
    def val(self, value):
        text = "" if value is None else str(value).strip()
        return text if text else NAO_LOCALIZADA

    def replace_supplier_intro(self, doc, ata):
        processo, pregao = self.extract_process_numbers(doc)
        intro = (
            "O CONSÓRCIO PÚBLICO INSTITUIÇÃO DE COOPERAÇÃO INTERMUNICIPAL DO MÉDIO PARAOPEBA - ICISMEP, "
            "CNPJ Nº 05.802.877/0001-10, órgão gerenciador, com sede na Rua Orquídeas, nº 489, Bairro Flor de Minas, "
            "no Município de São Joaquim de Bicas, Estado de Minas Gerais, CEP 32.920-000, a seguir denominado Consórcio ICISMEP, "
            f"neste ato representado por seu diretor institucional Sr. Eustáquio da Abadia Amaral e {self.val(ata.fornecedor_nome)}, "
            f"com sede/endereço {self.val(ata.endereco)}, Município {self.val(ata.municipio)}/{self.val(ata.uf)}, "
            f"CEP {self.val(ata.cep)}, telefone {self.val(ata.telefone)}, e-mail {self.val(ata.email)}, "
            f"inscrita no CNPJ sob o nº {self.val(ata.fornecedor_cnpj)}, Inscrição Estadual nº {self.val(ata.inscricao_estadual)}, "
            f"neste ato representada por seu representante legal {self.val(ata.representante)}, CPF nº {self.val(ata.cpf_representante)}, "
            f"RG nº {self.val(ata.rg_representante)}, órgão expedidor {self.val(ata.orgao_expedidor)}, "
            f"observadas, ainda, as disposições do Edital do PROCESSO LICITATÓRIO Nº {processo}, "
            f"na modalidade PREGÃO ELETRÔNICO Nº {pregao}, resolvem registrar os preços da empresa acima citada, "
            "de acordo com o item disputado e a classificação por ela alcançada, observadas as condições do Edital que integram "
            "este instrumento de registro, mediante as condições a seguir situadas:"
        )

        for p in doc.paragraphs:
            n = self.norm(p.text)
            if "CONSORCIO PUBLICO" in n and "RESOLVEM REGISTRAR" in n:
                self.set_paragraph_text(p, intro)
                return

    def extract_process_numbers(self, doc):
        text = "\n".join(p.text for p in doc.paragraphs[:20])
        processo = self.find_re(text, r"PROCESSO LICITAT[ÓO]RIO\s*N[º°]?\s*([0-9]+/20[0-9]{2})", "53/2026")
        pregao = self.find_re(text, r"PREG[ÃA]O ELETR[ÔO]NICO\s*N[º°]?\s*([0-9]+/20[0-9]{2})", "41/2026")
        return processo, pregao

    def find_re(self, text, pattern, default):
        m = re.search(pattern, text, flags=re.IGNORECASE)
        return m.group(1) if m else default

    # ------------------------ tabela cláusula 4 ------------------------
    def replace_registered_prices_table(self, doc, ata):
        table = self.find_price_table(doc)
        if table is None or len(table.rows) < 2:
            return

        header_idx = 0
        data_template = deepcopy(table.rows[1]._tr)
        total_template = deepcopy(table.rows[-1]._tr)

        while len(table.rows) > 1:
            tr = table.rows[-1]._tr
            tr.getparent().remove(tr)

        for item in ata.itens:
            table._tbl.append(deepcopy(data_template))
            row = table.rows[-1]
            self.clear_row(row)
            values = [
                item.codigo_siplan,
                str(item.numero_item),
                self.format_qty(item.quantidade),
                item.descricao_oficial,
                item.apresentacao,
                item.marca or item.fabricante or item.modelo,
                self.format_money(item.valor_unitario, 4),
                self.format_money(item.valor_total, 2),
            ]
            for i, value in enumerate(values):
                if i < len(row.cells):
                    self.set_cell_text(row.cells[i], value)

        table._tbl.append(deepcopy(total_template))
        total_row = table.rows[-1]
        self.clear_row(total_row)
        self.write_total_row(total_row, ata.valor_total)

    def find_price_table(self, doc):
        for table in doc.tables:
            if not table.rows:
                continue
            header = self.norm(" ".join(c.text for c in table.rows[0].cells))
            if "SIPLAN" in header and "DESCRI" in header and ("PRECO" in header or "VALOR" in header):
                return table
        return None

    def write_total_row(self, row, total):
        cells = row.cells
        if not cells:
            return
        try:
            label = cells[0]
            for idx in range(1, len(cells) - 1):
                label = label.merge(cells[idx])
            self.set_cell_text(label, "VALOR TOTAL:")
            self.set_cell_text(cells[-1], self.format_money(total, 2))
            self.make_horizontal(label)
            self.center(label)
            self.center(cells[-1])
            self.bold(label)
            self.bold(cells[-1])
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Cm(0.55)
        except Exception:
            self.set_cell_text(cells[0], f"VALOR TOTAL: {self.format_money(total, 2)}")

    # ------------------------ parágrafo total ------------------------
    def replace_total_paragraph(self, doc, ata):
        valor = self.format_money(ata.valor_total, 2)
        try:
            extenso = valor_por_extenso(ata.valor_total)
        except Exception:
            extenso = ""
        body = f"Valor total dos preços registrados: {valor} ({extenso})."

        for p in doc.paragraphs:
            if "Valor total dos preços registrados" in p.text:
                novo = body if self.paragraph_has_numbering(p) else f"4.2 {body}"
                self.set_paragraph_text(p, novo)
                return

    def paragraph_has_numbering(self, p):
        ppr = p._p.pPr
        if ppr is not None and ppr.numPr is not None:
            return True

        style = getattr(p, "style", None)
        while style is not None:
            element = getattr(style, "_element", None)
            style_ppr = getattr(element, "pPr", None)
            if style_ppr is not None and style_ppr.numPr is not None:
                return True
            style = getattr(style, "base_style", None)

        return False

    # ------------------------ apêndice ------------------------
    def replace_appendix_table(self, doc, ata):
        table = self.find_appendix_table(doc)
        if table is None or len(table.rows) < 2:
            return

        mapping = self.map_appendix_columns(table)
        data_template = deepcopy(table.rows[1]._tr)
        while len(table.rows) > 1:
            tr = table.rows[-1]._tr
            tr.getparent().remove(tr)

        for item in ata.itens:
            table._tbl.append(deepcopy(data_template))
            row = table.rows[-1]
            self.clear_row(row)

            values = item.appendix_cells_text or []
            if values and len(values) >= len(row.cells):
                self.write_appendix_values(row, values)
            else:
                self.write_appendix_item(row, mapping, item)

    def write_appendix_values(self, row, values):
        for i, value in enumerate(values):
            if i < len(row.cells):
                self.set_cell_text(row.cells[i], value)

    def write_appendix_item(self, row, mapping, item):
        data = {
            "codigo": item.codigo_siplan,
            "item": str(item.numero_item),
            "descricao": item.descricao_oficial,
            "apresentacao": item.apresentacao,
            "total": self.format_qty_total(item.quantidade),
        }
        for key, value in data.items():
            idx = mapping.get(key)
            if idx is not None and idx < len(row.cells):
                self.set_cell_text(row.cells[idx], value)

    def map_appendix_columns(self, table):
        if not table.rows:
            return {}

        headers = [self.norm(c.text) for c in table.rows[0].cells]
        mapping = {}
        for idx, header in enumerate(headers):
            if not header:
                continue
            if "codigo" not in mapping and ("SIPLAN" in header or "CODIGO" in header or header == "COD"):
                mapping["codigo"] = idx
            elif "item" not in mapping and re.search(r"\bITEM\b", header):
                mapping["item"] = idx
            elif "descricao" not in mapping and ("DESCRIT" in header or "DESCRICAO" in header):
                mapping["descricao"] = idx
            elif "apresentacao" not in mapping and ("APRESENT" in header or "UNIDADE" in header or header == "UN"):
                mapping["apresentacao"] = idx
            elif "total" not in mapping and (
                "DEMANDA TOTAL" in header
                or "TOTAL ENTES" in header
                or "CONSORCIADOS" in header
                or header == "TOTAL"
            ):
                mapping["total"] = idx

        if "total" not in mapping and headers:
            mapping["total"] = len(headers) - 1
        return mapping

    def find_appendix_table(self, doc):
        candidates = []
        for table in doc.tables:
            if not table.rows:
                continue
            header = self.norm(" ".join(c.text for c in table.rows[0].cells))
            score = 0
            for word in ["SIPLAN", "ITEM", "DESCRITIVO", "APRESENT", "ICISMEP", "TOTAL"]:
                if word in header:
                    score += 1
            if score >= 4 and len(table.rows[0].cells) > 10:
                candidates.append((score, table))
        if not candidates:
            return None
        return sorted(candidates, key=lambda x: x[0], reverse=True)[0][1]

    # ------------------------ assinatura ------------------------
    def replace_signature(self, doc, ata):
        fornecedor = ata.fornecedor_nome.title()
        representante = self.val(ata.representante)
        text = f"{representante}\n{fornecedor}"

        if len(doc.tables) >= 3:
            table = doc.tables[2]
            try:
                cell = table.rows[0].cells[-1]
                self.set_cell_text(cell, text)
            except Exception:
                pass

    # ------------------------ utilitários ------------------------
    def clear_row(self, row):
        for cell in row.cells:
            self.set_cell_text(cell, "")

    def set_cell_text(self, cell, text):
        text = "" if text is None else str(text)
        if cell.paragraphs:
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].text = text
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(text)
            for p2 in cell.paragraphs[1:]:
                for r in p2.runs:
                    r.text = ""
        else:
            cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(8)

    def set_paragraph_text(self, paragraph, text):
        if paragraph.runs:
            paragraph.runs[0].text = text
            for r in paragraph.runs[1:]:
                r.text = ""
        else:
            paragraph.add_run(text)

    def format_money(self, value, casas):
        value = float(value or 0)
        s = f"R$ {value:,.{casas}f}"
        return s.replace(",", "X").replace(".", ",").replace("X", ".")

    def format_qty(self, value):
        value = float(value or 0)
        if value.is_integer():
            return str(int(value))
        return str(value).replace(".", ",")

    def format_qty_total(self, value):
        value = float(value or 0)
        if value.is_integer():
            return f"{int(round(value)):,}".replace(",", ".")
        return str(value).replace(".", ",")

    def make_horizontal(self, cell):
        tcPr = cell._tc.get_or_add_tcPr()
        for child in list(tcPr):
            if child.tag == qn("w:textDirection"):
                tcPr.remove(child)
        text_dir = OxmlElement("w:textDirection")
        text_dir.set(qn("w:val"), "lrTb")
        tcPr.append(text_dir)

    def center(self, cell):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def bold(self, cell):
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    def norm(self, text):
        text = (text or "").upper()
        repl = {"Á":"A","À":"A","Ã":"A","Â":"A","É":"E","Ê":"E","Í":"I","Ó":"O","Õ":"O","Ô":"O","Ú":"U","Ç":"C"}
        for a,b in repl.items():
            text = text.replace(a,b)
        return " ".join(text.split())
