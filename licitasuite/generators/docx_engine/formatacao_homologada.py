from __future__ import annotations

from pathlib import Path
from typing import Any
import re
from zipfile import ZipFile, ZIP_DEFLATED
import xml.etree.ElementTree as ET

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


NAO_LOCALIZADA = "[INFORMAÇÃO NÃO LOCALIZADA]"


def _texto(valor: Any) -> str:
    return str(valor or "").strip()


def _formatar_milhar(valor: Any) -> str:
    texto = _texto(valor)
    if not texto:
        return ""

    texto_limpo = texto.replace(".", "").replace(",", "").strip()

    if not re.fullmatch(r"\d+", texto_limpo):
        return texto

    try:
        return f"{int(texto_limpo):,}".replace(",", ".")
    except Exception:
        return texto


def _title_empresa(nome: str) -> str:
    especiais = {"LTDA", "S.A.", "SA", "ME", "EPP", "EIRELI"}
    partes = _texto(nome).lower().split()
    saida = []

    for parte in partes:
        p = parte.upper().replace("S/A", "S.A.")
        saida.append(p if p in especiais else parte.capitalize())

    return " ".join(saida)


def _cell_text(cell) -> str:
    return " ".join(p.text for p in cell.paragraphs).strip()


def _limpar_paragrafo(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _run(paragraph, text: str, bold: bool = False, size: int = 11):
    r = paragraph.add_run(text)
    r.bold = bool(bold)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    return r


def _set_run(run, bold: bool = False, size: int = 11):
    run.bold = bool(bold)
    run.font.name = "Arial"
    run.font.size = Pt(size)


def _set_cell_text(cell, text: str, size: int = 8, bold: bool = False, center: bool = True):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_run(r, bold=bold, size=size)


def _norm_key(name: str) -> str:
    text = str(name or "").strip().lower()
    repl = {
        "á": "a", "à": "a", "â": "a", "ã": "a",
        "é": "e", "ê": "e",
        "í": "i",
        "ó": "o", "ô": "o", "õ": "o",
        "ú": "u",
        "ç": "c",
    }
    for a, b in repl.items():
        text = text.replace(a, b)
    text = re.sub(r"[^a-z0-9]+", "_", text)
    return text.strip("_")


def _norm_match(value: Any) -> str:
    text = str(value or "").upper()
    repl = {
        "Á": "A", "À": "A", "Â": "A", "Ã": "A",
        "É": "E", "Ê": "E", "Í": "I",
        "Ó": "O", "Ô": "O", "Õ": "O",
        "Ú": "U", "Ç": "C",
    }
    for a, b in repl.items():
        text = text.replace(a, b)
    text = re.sub(r"[^A-Z0-9]+", " ", text)
    tokens = text.split()
    stop = {
        "LTDA", "EIRELI", "SA", "S", "A", "ME", "EPP", "SS",
        "COMERCIO", "COM", "DISTRIBUIDORA", "DISTRIBUICAO",
        "MEDICAMENTOS", "PRODUTOS", "PRODUTO", "HOSPITALARES",
        "IMPORTACAO", "EXPORTACAO", "SERVICOS", "FARMACEUTICA",
        "FARMACEUTICOS",
    }
    return " ".join(t for t in tokens if t not in stop)


def _xlsx_col_to_index(ref: str) -> int:
    letters = re.sub(r"[^A-Z]", "", str(ref or "").upper())
    col = 0
    for ch in letters:
        col = col * 26 + (ord(ch) - 64)
    return col - 1


def _read_banco_xlsx(path: str | Path | None) -> list[dict[str, str]]:
    if not path:
        return []

    path = Path(path)
    if not path.exists():
        return []

    try:
        with ZipFile(path, "r") as z:
            names = z.namelist()

            shared = []
            if "xl/sharedStrings.xml" in names:
                root = ET.fromstring(z.read("xl/sharedStrings.xml"))
                ns = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
                for si in root.findall("a:si", ns):
                    shared.append("".join(t.text or "" for t in si.findall(".//a:t", ns)))

            sheet_name = "xl/worksheets/sheet1.xml"
            if sheet_name not in names:
                sheets = [n for n in names if n.startswith("xl/worksheets/sheet") and n.endswith(".xml")]
                if not sheets:
                    return []
                sheet_name = sheets[0]

            root = ET.fromstring(z.read(sheet_name))
            ns = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}

            matrix = []
            for row in root.findall(".//a:row", ns):
                values = {}
                max_col = -1
                for c in row.findall("a:c", ns):
                    idx = _xlsx_col_to_index(c.attrib.get("r", ""))
                    max_col = max(max_col, idx)
                    v = c.find("a:v", ns)
                    if v is None:
                        value = ""
                    elif c.attrib.get("t") == "s":
                        value = shared[int(v.text)] if v.text and int(v.text) < len(shared) else ""
                    else:
                        value = v.text or ""
                    values[idx] = str(value).strip()

                if max_col >= 0:
                    matrix.append([values.get(i, "") for i in range(max_col + 1)])

            if not matrix:
                return []

            headers = [_norm_key(h) for h in matrix[0]]
            rows = []
            for raw in matrix[1:]:
                row = {}
                for idx, header in enumerate(headers):
                    if header:
                        row[header] = raw[idx].strip() if idx < len(raw) else ""
                if any(row.values()):
                    rows.append(row)
            return rows
    except Exception:
        return []


def _cadastro_get(cadastro: dict[str, str] | None, *names: str) -> str:
    if not cadastro:
        return ""

    for name in names:
        key = _norm_key(name)
        value = cadastro.get(key, "")
        if value:
            return _texto(value)

    return ""


def _find_cadastro_for_ata(rows: list[dict[str, str]], ata: Any) -> dict[str, str] | None:
    if not rows:
        return None

    nome_ata = _ata_empresa(ata)
    cnpj_ata = _ata_cnpj(ata)
    nome_key = _norm_match(nome_ata)

    if cnpj_ata and cnpj_ata != NAO_LOCALIZADA:
        digits = re.sub(r"\D+", "", cnpj_ata)
        for row in rows:
            cnpj = row.get("cnpj", "")
            if digits and re.sub(r"\D+", "", cnpj) == digits:
                return row

    for row in rows:
        fornecedor = row.get("fornecedor", "") or row.get("razao_social", "") or row.get("nome", "")
        fornecedor_key = _norm_match(fornecedor)
        if not fornecedor_key or not nome_key:
            continue

        if fornecedor_key == nome_key or fornecedor_key in nome_key or nome_key in fornecedor_key:
            return row

        a = set(nome_key.split())
        b = set(fornecedor_key.split())
        if a and b and len(a & b) >= min(2, len(a), len(b)):
            return row

    return None


def _dict_like_get(obj: Any, *names: str) -> str:
    if isinstance(obj, dict):
        norm_map = {_norm_key(k): v for k, v in obj.items()}
        for name in names:
            value = norm_map.get(_norm_key(name))
            if value:
                return _texto(value)
    return ""


def _attr_get(obj: Any, *names: str) -> str:
    if obj is None:
        return ""

    # dict ou similar
    val = _dict_like_get(obj, *names)
    if val:
        return val

    # atributos comuns
    for name in names:
        candidates = {
            name,
            name.lower(),
            name.upper(),
            _norm_key(name),
        }

        for candidate in candidates:
            value = getattr(obj, candidate, None)
            if value:
                return _texto(value)

    # alguns objetos usam __dict__
    data = getattr(obj, "__dict__", None)
    if isinstance(data, dict):
        val = _dict_like_get(data, *names)
        if val:
            return val

    return ""


def _safe_get(obj: Any, *names: str, default: str = "") -> str:
    value = _attr_get(obj, *names)
    return value if value else default


def _get_from_supplier(ata: Any, *names: str, default: str = "") -> str:
    for source_name in ("fornecedor", "supplier", "dados_fornecedor", "fornecedor_dados", "cadastro", "dados"):
        source = getattr(ata, source_name, None)
        value = _attr_get(source, *names)
        if value:
            return value

    # também tenta na própria ata por segurança
    value = _attr_get(ata, *names)
    return value if value else default


def _fallback(value: str) -> str:
    return value if _texto(value) else NAO_LOCALIZADA


def _ata_empresa(ata: Any) -> str:
    return (
        _safe_get(ata, "fornecedor_nome", "razao_social", "razão social", "nome_fornecedor", "FORNECEDOR")
        or _get_from_supplier(ata, "FORNECEDOR", "fornecedor", "razao_social", "razão social", "nome", "fornecedor_nome")
    )


def _ata_endereco(ata: Any) -> str:
    return _fallback(_safe_get(ata, "endereco_fornecedor", "endereco", "ENDEREÇO") or _get_from_supplier(ata, "ENDEREÇO", "endereco"))


def _ata_cep(ata: Any) -> str:
    return _fallback(_safe_get(ata, "cep", "CEP") or _get_from_supplier(ata, "CEP", "cep"))


def _ata_fone(ata: Any) -> str:
    return _fallback(_safe_get(ata, "telefone", "fone", "FONE") or _get_from_supplier(ata, "FONE", "telefone", "fone"))


def _ata_email(ata: Any) -> str:
    return _fallback(_safe_get(ata, "email", "e_mail", "EMAIL") or _get_from_supplier(ata, "EMAIL", "email", "e_mail"))


def _ata_cnpj(ata: Any) -> str:
    return _fallback(_safe_get(ata, "cnpj", "CNPJ") or _get_from_supplier(ata, "CNPJ", "cnpj"))


def _ata_ie(ata: Any) -> str:
    return (
        _safe_get(ata, "inscricao_estadual", "inscrição estadual", "INSCRIÇÃO ESTAUDAL", "INSCRIÇÃO ESTADUAL", "ie")
        or _get_from_supplier(ata, "INSCRIÇÃO ESTAUDAL", "INSCRIÇÃO ESTADUAL", "inscricao_estadual", "inscrição estadual", "ie")
    )


def _ata_representante(ata: Any) -> str:
    return _fallback(
        _safe_get(ata, "representante", "representante_legal", "REPRESENTANTE", "socio")
        or _get_from_supplier(ata, "REPRESENTANTE", "representante", "representante_legal", "socio")
    )


def _ata_cpf(ata: Any) -> str:
    return _fallback(_safe_get(ata, "cpf", "CPF") or _get_from_supplier(ata, "CPF", "cpf"))


def _ata_rg(ata: Any) -> str:
    return _fallback(_safe_get(ata, "rg", "RG") or _get_from_supplier(ata, "RG", "rg"))


def _ata_orgao(ata: Any) -> str:
    return _fallback(
        _safe_get(ata, "orgao", "órgão", "orgao_expedidor", "ORGAO", "ÓRGÃO")
        or _get_from_supplier(ata, "ORGAO", "ÓRGÃO", "orgao", "órgão", "orgao_expedidor")
    )


def _extrair_processo_pregao(texto: str) -> tuple[str, str]:
    processo = ""
    pregao = ""

    m = re.search(r"PROCESSO\s+LICITAT[ÓO]RIO\s+N[º°]\s*([\d./-]+)", texto, flags=re.I)
    if m:
        processo = m.group(1)

    m = re.search(r"PREG[ÃA]O\s+ELETR[ÔO]NICO\s+N[º°]\s*([\d./-]+)", texto, flags=re.I)
    if m:
        pregao = m.group(1)

    return processo, pregao


def _normalizar_representante_assinatura(representante: str) -> str:
    rep = _texto(representante)
    rep = re.sub(r"^(seu|sua)\s+s[oó]cio\s+Sr\.?\s+", "", rep, flags=re.I).strip()
    rep = re.sub(r"^(Sr\.?|Sra\.?)\s+", "", rep, flags=re.I).strip()
    return rep or representante


def formatar_preambulo(document: Document, ata: Any, cadastro: dict[str, str] | None = None):
    empresa = _ata_empresa(ata).upper()
    if not empresa:
        return

    for p in document.paragraphs:
        txt = p.text or ""
        txt_norm = txt.upper()

        if "NESTE ATO REPRESENTADO POR SEU DIRETOR INSTITUCIONAL" not in txt_norm:
            continue

        if "PROCESSO LICITATÓRIO" not in txt_norm and "PROCESSO LICITATORIO" not in txt_norm:
            continue

        processo, pregao = _extrair_processo_pregao(txt)

        endereco = _fallback(_cadastro_get(cadastro, "ENDEREÇO", "ENDERECO") or _ata_endereco(ata))
        cep = _fallback(_cadastro_get(cadastro, "CEP") or _ata_cep(ata))
        fone = _fallback(_cadastro_get(cadastro, "FONE", "TELEFONE") or _ata_fone(ata))
        email = _fallback(_cadastro_get(cadastro, "EMAIL", "E-MAIL") or _ata_email(ata))
        cnpj = _fallback(_cadastro_get(cadastro, "CNPJ") or _ata_cnpj(ata))
        ie = _cadastro_get(cadastro, "INSCRIÇÃO ESTAUDAL", "INSCRIÇÃO ESTADUAL", "IE") or _ata_ie(ata)
        representante = _fallback(_cadastro_get(cadastro, "REPRESENTANTE") or _ata_representante(ata))
        cpf = _fallback(_cadastro_get(cadastro, "CPF") or _ata_cpf(ata))
        rg = _fallback(_cadastro_get(cadastro, "RG") or _ata_rg(ata))
        orgao = _fallback(_cadastro_get(cadastro, "ORGAO", "ÓRGÃO", "ORGAO EXPEDIDOR", "ÓRGÃO EXPEDIDOR") or _ata_orgao(ata))

        _limpar_paragrafo(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        _run(p, "O ", False)
        _run(p, "CONSÓRCIO PÚBLICO INSTITUIÇÃO DE COOPERAÇÃO INTERMUNICIPAL DO MÉDIO PARAOPEBA - ICISMEP, CNPJ Nº 05.802.877/0001-10", True)
        _run(p, ", órgão gerenciador, com sede na Rua Marciano Henriques, n° 107, Bairro Centro, no Município de Igarapé, Estado de Minas Gerais, CEP 32.510-008, a seguir denominado Consórcio ICISMEP, neste ato representado por seu diretor institucional Sr. ", False)
        _run(p, "Eustáquio da Abadia Amaral", False)
        _run(p, " e ", False)
        _run(p, empresa, True)
        _run(p, f", com sede na {endereco}, CEP {cep}, Fone {fone}, e-mail {email}, inscrita no CNPJ sob o n.º {cnpj}", False)

        if ie:
            _run(p, f", Inscrição Estadual n.º {ie}", False)

        _run(p, ", neste ato representada por ", False)
        _run(p, representante, True)
        _run(p, f", inscrito no CPF sob o nº {cpf} e portador da Carteira de Identidade n° {rg}, expedida pela {orgao}, nos termos do artigo 40, II da Lei Federal n° 14.133/21, observadas, ainda, as disposições do Edital do ", False)
        _run(p, f"PROCESSO LICITATÓRIO Nº {processo}", True)
        _run(p, ", na modalidade ", False)
        _run(p, f"PREGÃO ELETRÔNICO Nº {pregao}", True)
        _run(p, ", do tipo menor preço, auxiliado pelo Sistema de Registro de Preços, regido pela Lei Federal n° 14.133/21, e regulamentado pelo Decreto Federal n° 11.462/23, e demais disposições legais aplicáveis, de acordo com o resultado da classificação das propostas apresentadas no Pregão, resolvem registrar os preços da empresa acima citada, de acordo com o item disputado e a classificação por ela alcançada, observadas as condições do Edital que integram este instrumento de registro, mediante as condições a seguir situadas:", False)
        return


def formatar_quantidade_clausula_4(document: Document):
    for table in document.tables:
        quant_idx = None

        for row in table.rows[:3]:
            headers = [_cell_text(c).upper().replace(".", "").strip() for c in row.cells]
            if "QUANT" in headers:
                quant_idx = headers.index("QUANT")
                break

        if quant_idx is None:
            continue

        for row in table.rows[1:]:
            if quant_idx >= len(row.cells):
                continue

            cell = row.cells[quant_idx]
            txt = _cell_text(cell)
            novo = _formatar_milhar(txt)

            if novo != txt or txt:
                _set_cell_text(cell, novo, size=8, bold=False, center=True)


def formatar_assinaturas(document: Document, ata: Any, cadastro: dict[str, str] | None = None):
    empresa = _cadastro_get(cadastro, "FORNECEDOR", "RAZÃO SOCIAL", "RAZAO SOCIAL") or _ata_empresa(ata)
    representante = _cadastro_get(cadastro, "REPRESENTANTE") or _ata_representante(ata)

    if not empresa or not representante:
        return

    empresa_ass = _title_empresa(empresa)
    representante_ass = _normalizar_representante_assinatura(representante)

    for p in document.paragraphs:
        txt = (p.text or "").strip()
        low = txt.lower()

        if (empresa.lower() in low or representante.lower() in low) and len(txt) <= 180:
            _limpar_paragrafo(p)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = _cell_text(cell)
                low = text.lower()

                if empresa.lower() in low or representante.lower() in low:
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    _run(p, "\n\n\n\n", False)
                    _run(p, representante_ass, True)
                    _run(p, "\n", False)
                    _run(p, empresa_ass, True)


def aplicar_formatacao_homologada(docx_path: str | Path, ata: Any, cadastro: dict[str, str] | None = None):
    docx_path = Path(docx_path)
    document = Document(docx_path)

    formatar_preambulo(document, ata, cadastro)
    formatar_quantidade_clausula_4(document)
    formatar_assinaturas(document, ata, cadastro)

    document.save(docx_path)


def aplicar_em_lote(files: list, atas: list, banco_path: str | Path | None = None):
    banco_rows = _read_banco_xlsx(banco_path)
    for idx, file in enumerate(files):
        ata = atas[idx] if idx < len(atas) else None
        if not ata:
            continue
        cadastro = _find_cadastro_for_ata(banco_rows, ata)
        aplicar_formatacao_homologada(file, ata, cadastro)


def recriar_zip(files: list, zip_path: str | Path):
    zip_path = Path(zip_path)
    if not files or not zip_path:
        return zip_path

    with ZipFile(zip_path, "w", ZIP_DEFLATED) as z:
        for f in files:
            p = Path(f)
            if p.exists():
                z.write(p, p.name)

    return zip_path
