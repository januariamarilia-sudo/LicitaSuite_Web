from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
from copy import deepcopy
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from licitasuite.utils.formatters import safe_filename, format_brl, format_quantity
from licitasuite.utils.value_extenso import valor_por_extenso


HEADER_BLUE = "A9BDD9"
TABLE_FONT = "Arial"
TABLE_FONT_SIZE = 8
BLANK = "____________________________"


class AtaGenerator:
    def __init__(self, template_path, output_dir="output/atas_geradas"):
        self.template_path = Path(template_path)
        self.output_dir = Path(output_dir)

    def generate_all(self, atas):
        self.output_dir.mkdir(parents=True, exist_ok=True)

        for old in self.output_dir.glob("*.docx"):
            try:
                old.unlink()
            except Exception:
                pass

        generated = []
        for ata in atas:
            if ata.itens:
                generated.append(self.generate_one(ata))

        zip_path = self.output_dir.parent / "atas_geradas.zip"
        if zip_path.exists():
            zip_path.unlink()

        with ZipFile(zip_path, "w", ZIP_DEFLATED) as z:
            for file in generated:
                z.write(file, file.name)

        return generated, zip_path

    def generate_one(self, ata):
        doc = Document(self.template_path)

        self.replace_supplier_placeholders(doc, ata)
        self.replace_intro_supplier_paragraph(doc, ata)
        self.replace_signature_block(doc, ata)

        self.rebuild_registered_prices_table(doc, ata)
        self.rebuild_appendix_with_original_rows(doc, ata)
        self.fix_total_value_cells(doc, ata)
        self.fix_appendix_last_header(doc)
        self.apply_table_visual_standard(doc)

        filename = f"ATA DE REGISTRO DE PREÇOS - {safe_filename(ata.fornecedor_nome)}.docx"
        output = self.output_dir / filename
        doc.save(output)
        return output

    # =========================================================
    # DADOS DO FORNECEDOR
    # =========================================================
    def clean(self, value):
        text = "" if value is None else str(value).strip()
        if text.lower() in ["nan", "none", "null", ""]:
            return BLANK
        return text

    def title_case_company(self, value):
        text = self.clean(value)
        if text == BLANK:
            return text
        keep_upper = {"LTDA", "EPP", "ME", "S/A", "SA", "EIRELI", "CNPJ", "CPF", "SS", "LC123"}
        parts = []
        for word in text.split():
            raw = word.strip()
            token = re.sub(r"[^A-Za-zÀ-ÿ0-9/.-]", "", raw).upper()
            if token in keep_upper or any(ch.isdigit() for ch in raw):
                parts.append(raw.upper())
            else:
                parts.append(raw[:1].upper() + raw[1:].lower())
        return " ".join(parts)

    def supplier_values(self, ata):
        return {
            "nome": self.clean(getattr(ata, "fornecedor_nome", "")),
            "nome_titulo": self.title_case_company(getattr(ata, "fornecedor_nome", "")),
            "cnpj": self.clean(getattr(ata, "fornecedor_cnpj", "")),
            "endereco": self.clean(getattr(ata, "endereco", "")),
            "municipio": self.clean(getattr(ata, "municipio", "")),
            "uf": self.clean(getattr(ata, "uf", "")),
            "cep": self.clean(getattr(ata, "cep", "")),
            "telefone": self.clean(getattr(ata, "telefone", "")),
            "email": self.clean(getattr(ata, "email", "")),
            "representante": self.title_case_company(getattr(ata, "representante", "")),
            "cpf": self.clean(getattr(ata, "cpf_representante", "")),
            "rg": self.clean(getattr(ata, "rg_representante", "")),
            "valor_total": format_brl(getattr(ata, "valor_total", 0), 2),
            "valor_total_extenso": valor_por_extenso(getattr(ata, "valor_total", 0)),
        }

    def supplier_map(self, ata):
        v = self.supplier_values(ata)
        return {
            "{{FORNECEDOR}}": v["nome"],
            "{{RAZAO_SOCIAL}}": v["nome"],
            "{{RAZÃO_SOCIAL}}": v["nome"],
            "{{CNPJ}}": v["cnpj"],
            "{{ENDERECO}}": v["endereco"],
            "{{ENDEREÇO}}": v["endereco"],
            "{{MUNICIPIO}}": v["municipio"],
            "{{MUNICÍPIO}}": v["municipio"],
            "{{UF}}": v["uf"],
            "{{CEP}}": v["cep"],
            "{{TELEFONE}}": v["telefone"],
            "{{EMAIL}}": v["email"],
            "{{E-MAIL}}": v["email"],
            "{{REPRESENTANTE}}": v["representante"],
            "{{REPRESENTANTE_LEGAL}}": v["representante"],
            "{{CPF_REPRESENTANTE}}": v["cpf"],
            "{{RG_REPRESENTANTE}}": v["rg"],
            "{{VALOR_TOTAL}}": v["valor_total"],
            "{{VALOR_TOTAL_EXTENSO}}": v["valor_total_extenso"],
        }

    def replace_supplier_placeholders(self, doc, ata):
        replacements = self.supplier_map(ata)

        for paragraph in doc.paragraphs:
            self.replace_paragraph_text(paragraph, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_paragraph_text(paragraph, replacements)

        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                self.replace_paragraph_text(paragraph, replacements)
            for paragraph in section.footer.paragraphs:
                self.replace_paragraph_text(paragraph, replacements)

    def replace_intro_supplier_paragraph(self, doc, ata):
        v = self.supplier_values(ata)

        parts = [
            ("O ", False),
            ("CONSÓRCIO PÚBLICO INSTITUIÇÃO DE COOPERAÇÃO INTERMUNICIPAL DO MÉDIO PARAOPEBA - ICISMEP, CNPJ Nº 05.802.877/0001-10", True),
            (", órgão gerenciador, com sede na Rua Orquídeas, nº 489, Bairro Flor de Minas, no Município de São Joaquim de Bicas, Estado de Minas Gerais, CEP 32.920-000, a seguir denominado Consórcio ICISMEP, neste ato representado por seu diretor institucional Sr. Eustáquio da Abadia Amaral e ", False),
            (v["nome"], True),
            (", com sede/endereço ", False),
            (v["endereco"], False),
            (", Município ", False),
            (f"{v['municipio']}/{v['uf']}", False),
            (", CEP ", False),
            (v["cep"], False),
            (", telefone ", False),
            (v["telefone"], False),
            (", e-mail ", False),
            (v["email"], False),
            (", inscrita no CNPJ sob o nº ", False),
            (v["cnpj"], True),
            (", neste ato representada por seu representante legal ", False),
            (v["representante"], True),
            (", CPF nº ", False),
            (v["cpf"], False),
            (", RG nº ", False),
            (v["rg"], False),
            (", observadas, ainda, as disposições do Edital do ", False),
            ("PROCESSO LICITATÓRIO Nº 45/2026, na modalidade PREGÃO ELETRÔNICO Nº 35/2026", True),
            (", resolvem registrar os preços da empresa acima citada, de acordo com o item disputado e a classificação por ela alcançada, observadas as condições do Edital que integram este instrumento de registro, mediante as condições a seguir situadas:", False),
        ]

        for p in doc.paragraphs:
            t = self.norm(p.text)
            if ("CONSORCIO PUBLICO" in t and "CNPJ" in t and "RESOLVEM REGISTRAR" in t) or (
                "CENTRAL BRASIL" in t and "RESOLVEM REGISTRAR" in t
            ):
                self.set_paragraph_parts(p, parts)
                return

    def replace_signature_block(self, doc, ata):
        v = self.supplier_values(ata)
        fornecedor_signature = f"{v['nome_titulo']}\n{v['representante']}"

        for p in doc.paragraphs[-120:]:
            n = self.norm(p.text)
            if "CENTRAL BRASIL" in n or "LUCAS VINICIUS" in n:
                self.set_paragraph_text_keep_first_run(p, fornecedor_signature)

        for table in doc.tables[-8:]:
            for row in table.rows:
                for cell in row.cells:
                    n = self.norm(cell.text)
                    if "CENTRAL BRASIL" in n or "LUCAS VINICIUS" in n:
                        self.set_cell_text_keep_style(cell, fornecedor_signature)
                    elif "{{ASSINATURA_FORNECEDOR}}" in cell.text:
                        self.set_cell_text_keep_style(cell, fornecedor_signature)

    def replace_paragraph_text(self, paragraph, replacements):
        changed = False
        for run in paragraph.runs:
            original = run.text
            updated = original
            for old, new in replacements.items():
                updated = updated.replace(old, str(new))
            if updated != original:
                run.text = updated
                changed = True

        if changed:
            return

        full = "".join(run.text for run in paragraph.runs)
        updated = full
        for old, new in replacements.items():
            updated = updated.replace(old, str(new))

        if updated != full and paragraph.runs:
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ""

    def set_paragraph_parts(self, paragraph, parts):
        font_name = "Arial"
        font_size = None
        if paragraph.runs:
            font_name = paragraph.runs[0].font.name or "Arial"
            font_size = paragraph.runs[0].font.size
            for run in paragraph.runs:
                run.text = ""

        for text, bold in parts:
            run = paragraph.add_run(text)
            run.bold = bold
            run.font.name = font_name
            if font_size:
                run.font.size = font_size

    def set_paragraph_text_keep_first_run(self, paragraph, text):
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    # =========================================================
    # CLÁUSULA 4
    # =========================================================
    def rebuild_registered_prices_table(self, doc, ata):
        table = self.find_registered_prices_table(doc)
        if table is None:
            return

        header_index = self.find_price_header_row_index(table)
        if header_index is None:
            header_index = 0

        template_row_xml = self.find_price_template_row_xml(table, header_index)
        column_map = self.map_columns_from_header(table.rows[header_index])

        while len(table.rows) > header_index + 1:
            tr = table.rows[-1]._tr
            tr.getparent().remove(tr)

        for item in ata.itens:
            table._tbl.append(deepcopy(template_row_xml))
            row = table.rows[-1]
            self.clear_row_text(row)
            self.write_price_row(row, item, column_map)

        self.append_total_row(table, template_row_xml, ata)
        self.format_table(table, header_rows=[header_index])

    def append_total_row(self, table, template_row_xml, ata):
        table._tbl.append(deepcopy(template_row_xml))
        row = table.rows[-1]
        self.clear_row_text(row)
        self.set_total_row_compact_height(row)

        cells = row.cells
        if len(cells) >= 2:
            try:
                value_cell = cells[-1]
                label_cell = cells[0]

                # Mescla todas as células antes da coluna do valor,
                # mantendo o valor final separado, como no modelo.
                for idx in range(1, len(cells) - 1):
                    label_cell = label_cell.merge(cells[idx])

                self.set_cell_text_keep_style(label_cell, "VALOR TOTAL:")
                self.set_cell_text_keep_style(value_cell, format_brl(ata.valor_total, 2))

                # Ajuste final solicitado: texto horizontal, não vertical.
                self.make_cell_text_horizontal(label_cell)
                self.center_cell(label_cell)
                self.center_cell(value_cell)

                self.bold_cell(label_cell)
                self.bold_cell(value_cell)
            except Exception:
                label_idx = max(0, len(cells) - 2)
                value_idx = len(cells) - 1
                self.set_cell_text_keep_style(cells[label_idx], "VALOR TOTAL:")
                self.set_cell_text_keep_style(cells[value_idx], format_brl(ata.valor_total, 2))
                self.make_cell_text_horizontal(cells[label_idx])
                self.center_cell(cells[label_idx])
                self.center_cell(cells[value_idx])
                self.bold_cell(cells[label_idx])
                self.bold_cell(cells[value_idx])
        elif cells:
            self.set_cell_text_keep_style(cells[0], f"VALOR TOTAL: {format_brl(ata.valor_total, 2)}")
            self.make_cell_text_horizontal(cells[0])
            self.center_cell(cells[0])
            self.bold_cell(cells[0])

    def find_registered_prices_table(self, doc):
        best, best_score = None, 0
        for table in doc.tables:
            text = self.norm(" ".join(cell.text for row in table.rows for cell in row.cells))
            score = sum(1 for w in ["SIPLAN", "ITEM", "DESCRI", "MARCA", "MODELO", "PRECO", "TOTAL", "QUANT"] if w in text)
            if any(m in text for m in ["MUNICIP", "QUANTITATIVO PARA", "CONTAGEM", "BETIM", "IGARAPE", "OURO PRETO"]):
                score -= 3
            if score > best_score:
                best, best_score = table, score
        return best if best_score >= 5 else None

    def find_price_header_row_index(self, table):
        for idx, row in enumerate(table.rows):
            text = self.norm(" ".join(cell.text for cell in row.cells))
            if ("DESCRI" in text and ("PRECO" in text or "TOTAL" in text)) or ("SIPLAN" in text and "ITEM" in text):
                return idx
        return None

    def find_price_template_row_xml(self, table, header_index):
        return deepcopy(table.rows[header_index + 1]._tr) if len(table.rows) > header_index + 1 else deepcopy(table.rows[header_index]._tr)

    def map_columns_from_header(self, header_row):
        mapping = {}
        for idx, cell in enumerate(header_row.cells):
            text = self.norm(cell.text)
            if ("SIPLAN" in text or "COD" in text) and "codigo" not in mapping:
                mapping["codigo"] = idx
            elif "ITEM" in text and "item" not in mapping:
                mapping["item"] = idx
            elif "QUANT" in text and "quantidade" not in mapping:
                mapping["quantidade"] = idx
            elif "DESCRI" in text and "descricao" not in mapping:
                mapping["descricao"] = idx
            elif "APRESENT" in text and "apresentacao" not in mapping:
                mapping["apresentacao"] = idx
            elif "MARCA" in text and "marca" not in mapping:
                mapping["marca"] = idx
            elif "MODELO" in text and "modelo" not in mapping:
                mapping["modelo"] = idx
            elif ("UNIT" in text or "UNITARIO" in text) and "unitario" not in mapping:
                mapping["unitario"] = idx
            elif "TOTAL" in text and "total" not in mapping:
                mapping["total"] = idx

        count = len(header_row.cells)
        if count >= 9 and len(mapping) < 7:
            mapping = {
                "codigo": 0, "item": 1, "quantidade": 2, "descricao": 3,
                "apresentacao": 4, "marca": 5, "modelo": 6,
                "unitario": 7, "total": 8,
            }
        return mapping

    def write_price_row(self, row, item, mapping):
        values = {
            "codigo": item.codigo_siplan,
            "item": str(item.numero_item),
            "quantidade": format_quantity(item.quantidade),
            "descricao": item.descricao_oficial,
            "apresentacao": item.apresentacao,
            "marca": item.marca,
            "modelo": item.modelo,
            "unitario": format_brl(item.valor_unitario, 4),
            "total": format_brl(item.valor_total, 2),
        }

        for field, idx in mapping.items():
            if idx < len(row.cells):
                if field == "descricao":
                    self.write_description_from_appendix(row.cells[idx], item)
                else:
                    self.set_cell_text_keep_style(row.cells[idx], values.get(field, ""))

    def write_description_from_appendix(self, target_cell, item):
        try:
            if getattr(item, "appendix_row_xml", None) is not None:
                src_tc = item.appendix_row_xml.tc_lst[2]
                self.replace_cell_xml_content(target_cell, src_tc)
                return
        except Exception:
            pass
        self.set_cell_text_keep_style(target_cell, item.descricao_oficial)

    # =========================================================
    # APÊNDICE
    # =========================================================
    def rebuild_appendix_with_original_rows(self, doc, ata):
        table = self.find_appendix_table(doc)
        if table is None:
            return

        header_index = self.find_appendix_header_row_index(table)
        if header_index is None:
            header_index = 0

        while len(table.rows) > header_index + 1:
            tr = table.rows[-1]._tr
            tr.getparent().remove(tr)

        for item in ata.itens:
            if getattr(item, "appendix_row_xml", None) is not None:
                table._tbl.append(deepcopy(item.appendix_row_xml))
            else:
                row_xml = deepcopy(table.rows[header_index]._tr)
                table._tbl.append(row_xml)
                row = table.rows[-1]
                self.clear_row_text(row)
                self.write_appendix_fallback(row, item)

        self.fix_appendix_last_header_in_table(table)
        self.format_table(table, header_rows=list(range(0, header_index + 1)), appendix=True)

    def find_appendix_table(self, doc):
        best, best_score = None, 0
        for table in doc.tables:
            text = self.norm(" ".join(cell.text for row in table.rows for cell in row.cells))
            score = sum(1 for w in ["CODIGO DO SIPLAN", "DESCRITIVO", "APRESENTACAO", "QUANTITATIVO", "MUNICIP", "TOTAL", "ITEM"] if w in text)
            if len(table.columns) >= 8:
                score += 2
            if score > best_score:
                best, best_score = table, score
        return best if best_score >= 5 else None

    def find_appendix_header_row_index(self, table):
        for idx, row in enumerate(table.rows):
            text = self.norm(" ".join(cell.text for cell in row.cells))
            if "DESCRITIVO" in text and ("MUNICIP" in text or "QUANTITATIVO" in text or "TOTAL" in text):
                return idx
        return None

    def write_appendix_fallback(self, row, item):
        values = getattr(item, "appendix_cells_text", None)
        if not values:
            values = [
                item.codigo_siplan,
                str(item.numero_item),
                item.descricao_oficial,
                item.apresentacao,
                format_quantity(item.quantidade),
                format_brl(item.valor_total, 2),
            ]

        for idx, cell in enumerate(row.cells):
            self.set_cell_text_keep_style(cell, values[idx] if idx < len(values) else "")

    def fix_appendix_last_header(self, doc):
        table = self.find_appendix_table(doc)
        if table:
            self.fix_appendix_last_header_in_table(table)

    def fix_appendix_last_header_in_table(self, table):
        for row in table.rows[:2]:
            if row.cells:
                last = row.cells[-1]
                if not last.text.strip() or "____" in last.text or "—" in last.text:
                    self.set_cell_text_keep_style(last, "TOTAL")
                    return

    # =========================================================
    # TOTAL NO TEXTO
    # =========================================================
    def fix_total_value_cells(self, doc, ata):
        value = format_brl(ata.valor_total, 2)
        extenso = valor_por_extenso(ata.valor_total)

        for p in doc.paragraphs:
            if "Valor total dos preços registrados" in p.text or "VALOR TOTAL DOS PRECOS REGISTRADOS" in self.norm(p.text):
                self.set_paragraph_text_keep_first_run(
                    p,
                    f"Valor total dos preços registrados: {value} ({extenso})."
                )

    # =========================================================
    # FORMATAÇÃO
    # =========================================================
    def format_table(self, table, header_rows=None, appendix=False):
        header_rows = header_rows or []
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                self.set_cell_font(cell, TABLE_FONT, TABLE_FONT_SIZE)
                if r_idx in header_rows or (appendix and c_idx in [0, 1]):
                    self.set_cell_shading(cell, HEADER_BLUE)

    def apply_table_visual_standard(self, doc):
        for table in doc.tables:
            text = self.norm(" ".join(cell.text for row in table.rows for cell in row.cells))
            is_price = "DESCRI" in text and ("PRECO" in text or "TOTAL" in text) and "SIPLAN" in text
            is_appendix = "DESCRITIVO" in text and ("MUNICIP" in text or "QUANTITATIVO" in text or "CODIGO DO SIPLAN" in text)
            if is_price:
                hi = self.find_price_header_row_index(table)
                self.format_table(table, header_rows=[hi if hi is not None else 0])
            if is_appendix:
                hi = self.find_appendix_header_row_index(table)
                rows = list(range(0, hi + 1)) if hi is not None else [0]
                self.format_table(table, header_rows=rows, appendix=True)

    def set_cell_font(self, cell, font_name, font_size):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.rFonts
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.append(rFonts)
                rFonts.set(qn("w:ascii"), font_name)
                rFonts.set(qn("w:hAnsi"), font_name)
                rFonts.set(qn("w:cs"), font_name)

    def set_cell_shading(self, cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:fill"), fill)

    def bold_cell(self, cell):
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


    def center_cell(self, cell):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def set_total_row_compact_height(self, row):
        """
        Deixa a linha VALOR TOTAL compacta, sem herdar a altura da última linha de item.
        """
        try:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Cm(0.55)
        except Exception:
            pass

        try:
            trPr = row._tr.get_or_add_trPr()
            for child in list(trPr):
                if child.tag == qn("w:trHeight"):
                    trPr.remove(child)

            tr_height = OxmlElement("w:trHeight")
            tr_height.set(qn("w:val"), "310")
            tr_height.set(qn("w:hRule"), "exact")
            trPr.append(tr_height)
        except Exception:
            pass

    def make_cell_text_horizontal(self, cell):
        tcPr = cell._tc.get_or_add_tcPr()

        # Remove qualquer direção antiga, especialmente vertical.
        for child in list(tcPr):
            if child.tag == qn("w:textDirection"):
                tcPr.remove(child)

        # Define explicitamente texto horizontal: esquerda para direita, topo para baixo.
        text_dir = OxmlElement("w:textDirection")
        text_dir.set(qn("w:val"), "lrTb")
        tcPr.append(text_dir)

    # =========================================================
    # UTILITÁRIOS
    # =========================================================
    def replace_cell_xml_content(self, target_cell, source_tc):
        for p in list(target_cell._tc.p_lst):
            target_cell._tc.remove(p)
        for tbl in list(target_cell._tc.tbl_lst):
            target_cell._tc.remove(tbl)
        for child in source_tc.iterchildren():
            if child.tag.endswith('}p') or child.tag.endswith('}tbl'):
                target_cell._tc.append(deepcopy(child))
        self.set_cell_font(target_cell, TABLE_FONT, TABLE_FONT_SIZE)

    def clear_row_text(self, row):
        for cell in row.cells:
            self.set_cell_text_keep_style(cell, "")

    def set_cell_text_keep_style(self, cell, value):
        text = "" if value is None else str(value)
        if cell.paragraphs:
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].text = text
                for run in p.runs[1:]:
                    run.text = ""
            else:
                p.add_run(text)
            for extra in cell.paragraphs[1:]:
                for run in extra.runs:
                    run.text = ""
        else:
            cell.text = text

    def norm(self, text):
        text = (text or "").upper()
        replacements = {
            "Á": "A", "À": "A", "Ã": "A", "Â": "A",
            "É": "E", "Ê": "E",
            "Í": "I",
            "Ó": "O", "Õ": "O", "Ô": "O",
            "Ú": "U",
            "Ç": "C",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return " ".join(text.split())
