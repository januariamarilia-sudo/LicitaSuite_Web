from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
from copy import deepcopy
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from licitasuite.parsers.text_utils import normalize_text, safe_filename, format_money, format_qty

NA = "[INFORMAÇÃO NÃO LOCALIZADA]"

def valor_extenso(value):
    try:
        from num2words import num2words
        reais = int(value)
        cents = int(round((float(value) - reais) * 100))
        txt = num2words(reais, lang="pt_BR")
        txt += " real" if reais == 1 else " reais"
        if cents:
            txt += " e " + num2words(cents, lang="pt_BR") + (" centavo" if cents == 1 else " centavos")
        return txt
    except Exception:
        return ""

class CopyModelAtaGenerator:
    """
    LicitaSuite Web 3.1.2 LTS

    Base: 3.1.1 Rollback + Cor.

    Correções pontuais:
    - Mantém a tabela da cláusula 4 depois do parágrafo 4.1.
    - Restaura negritos principais no parágrafo inicial.
    - Não altera largura, altura ou estrutura da tabela.
    """

    def __init__(self, modelo_path, output_dir="output/atas_geradas"):
        self.modelo_path = Path(modelo_path)
        self.output_dir = Path(output_dir)

    def generate_all(self, atas):
        self.output_dir.mkdir(parents=True, exist_ok=True)
        for f in self.output_dir.glob("*.docx"):
            try:
                f.unlink()
            except Exception:
                pass

        generated = [self.generate_one(a) for a in atas if a.itens]

        zip_path = self.output_dir.parent / "atas_geradas.zip"
        if zip_path.exists():
            zip_path.unlink()

        with ZipFile(zip_path, "w", ZIP_DEFLATED) as z:
            for f in generated:
                z.write(f, f.name)

        return generated, zip_path

    def generate_one(self, ata):
        doc = Document(self.modelo_path)
        processo, pregao = self.extract_process(doc)

        self.replace_intro(doc, ata, processo, pregao)
        self.replace_price_table(doc, ata)
        self.ensure_price_table_after_41(doc)
        self.replace_total_paragraph(doc, ata)
        self.replace_appendix(doc, ata)
        self.replace_signature(doc, ata)

        fornecedor = self.clean_supplier_name(ata.fornecedor_nome)
        out = self.output_dir / f"ATA DE REGISTRO DE PREÇOS - {safe_filename(fornecedor)}.docx"
        doc.save(out)
        return out

    def clean_supplier_name(self, name):
        name = "" if name is None else str(name).strip()
        for marker in [" | Tipo:", " - LC123:", " - Documento", " Documento "]:
            if marker in name:
                name = name.split(marker, 1)[0]
        return " ".join(name.split()) or "FORNECEDOR_NAO_IDENTIFICADO"

    def extract_process(self, doc):
        text = "\n".join(p.text for p in doc.paragraphs[:30])
        processo = self._find(text, r"PROCESSO LICITAT[ÓO]RIO\s*N[º°]?\s*([0-9]+/20[0-9]{2})", "")
        pregao = self._find(text, r"PREG[ÃA]O ELETR[ÔO]NICO\s*N[º°]?\s*([0-9]+/20[0-9]{2})", "")
        return processo or NA, pregao or NA

    def _find(self, text, pattern, default):
        m = re.search(pattern, text, flags=re.I)
        return m.group(1) if m else default

    def val(self, value):
        t = "" if value is None else str(value).strip()
        return t if t else NA

    def replace_intro(self, doc, ata, processo, pregao):
        fornecedor = self.clean_supplier_name(ata.fornecedor_nome)
        text = (
            "O CONSÓRCIO PÚBLICO INSTITUIÇÃO DE COOPERAÇÃO INTERMUNICIPAL DO MÉDIO PARAOPEBA - ICISMEP, "
            "CNPJ Nº 05.802.877/0001-10, órgão gerenciador, com sede na Rua Orquídeas, nº 489, Bairro Flor de Minas, "
            "no Município de São Joaquim de Bicas, Estado de Minas Gerais, CEP 32.920-000, a seguir denominado Consórcio ICISMEP, "
            f"neste ato representado por seu diretor institucional Sr. Eustáquio da Abadia Amaral e {fornecedor}, "
            f"com sede/endereço {self.val(ata.endereco)}, Município {self.val(ata.municipio)}/{self.val(ata.uf)}, "
            f"CEP {self.val(ata.cep)}, telefone {self.val(ata.telefone)}, e-mail {self.val(ata.email)}, "
            f"inscrita no CNPJ sob o nº {self.val(ata.fornecedor_cnpj)}, Inscrição Estadual nº {self.val(ata.inscricao_estadual)}, "
            f"neste ato representada por seu representante legal {self.val(ata.representante)}, CPF nº {self.val(ata.cpf_representante)}, "
            f"RG nº {self.val(ata.rg_representante)}, órgão expedidor {self.val(ata.orgao_expedidor)}, "
            f"observadas, ainda, as disposições do Edital do PROCESSO LICITATÓRIO Nº {processo}, "
            f"na modalidade PREGÃO ELETRÔNICO Nº {pregao}, resolvem registrar os preços da empresa acima citada, "
            "de acordo com o item disputado e a classificação por ela alcançada, observadas as condições do Edital que integram este instrumento de registro, "
            "mediante as condições a seguir situadas:"
        )
        for p in doc.paragraphs:
            n = normalize_text(p.text)
            if "CONSORCIO PUBLICO" in n and "RESOLVEM REGISTRAR" in n:
                self.set_intro_paragraph_with_bold(p, text, fornecedor, processo, pregao)
                return

    def set_intro_paragraph_with_bold(self, p, text, fornecedor, processo, pregao):
        """
        Recria apenas os runs do parágrafo inicial para restaurar os negritos principais.
        Não mexe no restante do documento.
        """
        bold_terms = [
            "O CONSÓRCIO PÚBLICO INSTITUIÇÃO DE COOPERAÇÃO INTERMUNICIPAL DO MÉDIO PARAOPEBA - ICISMEP",
            "Eustáquio da Abadia Amaral",
            fornecedor,
            f"PROCESSO LICITATÓRIO Nº {processo}",
            f"PREGÃO ELETRÔNICO Nº {pregao}",
        ]

        # captura estilo base do primeiro run
        base_run = p.runs[0] if p.runs else None
        for r in list(p.runs):
            r.text = ""

        pos = 0
        pattern = "(" + "|".join(re.escape(t) for t in bold_terms if t and t != NA) + ")"
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue
            run = p.add_run(part)
            if base_run is not None:
                run.font.name = base_run.font.name
                run.font.size = base_run.font.size
            run.bold = part in bold_terms

    # ------------------- TABELA CLÁUSULA 4 -------------------
    def replace_price_table(self, doc, ata):
        table = self.find_price_table(doc)
        if table is None:
            return

        mapping = self.map_header_columns(table.rows[0])
        data_idx = 1 if len(table.rows) > 1 else 0
        data_template = deepcopy(table.rows[data_idx]._tr)

        while len(table.rows) > 1:
            table.rows[-1]._tr.getparent().remove(table.rows[-1]._tr)

        for item in ata.itens:
            table._tbl.append(deepcopy(data_template))
            row = table.rows[-1]
            self.clear_row_text_only(row)

            self.put(row, mapping, "lote", getattr(item, "lote_display", None) or getattr(item, "lote", "") or item.codigo_siplan)
            self.put(row, mapping, "siplan", item.codigo_siplan)
            self.put(row, mapping, "item", getattr(item, "item_display", None) or str(item.numero_item))
            self.put(row, mapping, "quant", format_qty(item.quantidade, use_thousands=False))
            self.put(row, mapping, "descricao", item.descricao_oficial, description=True)
            self.put(row, mapping, "apresentacao", item.apresentacao)

            marca_text = item.marca or item.fabricante or ""
            modelo_text = item.modelo or ""

            if "modelo" not in mapping and modelo_text:
                marca_text = (marca_text + " " + modelo_text).strip()

            self.put(row, mapping, "marca", marca_text)
            self.put(row, mapping, "modelo", modelo_text)
            self.put(row, mapping, "preco_unitario", format_money(item.valor_unitario, 4))
            self.put(row, mapping, "preco_total", format_money(item.valor_total, 2))

        table._tbl.append(deepcopy(data_template))
        total_row = table.rows[-1]
        self.clear_row_text_only(total_row)
        self.write_total(total_row, mapping, ata.valor_total)

    def ensure_price_table_after_41(self, doc):
        """
        Corrige somente a posição:
        no modelo correto, a ordem é:
        4 DOS PREÇOS REGISTRADOS
        4.1 Os preços...
        TABELA
        4.2 Valor total...
        """
        table = self.find_price_table(doc)
        if table is None:
            return

        p41 = None
        for p in doc.paragraphs:
            if normalize_text(p.text).startswith("4.1 OS PRECOS REGISTRADOS"):
                p41 = p
                break
        if p41 is None:
            return

        # Move o XML da tabela para imediatamente depois do parágrafo 4.1.
        tbl = table._tbl
        parent = tbl.getparent()
        if parent is None:
            return
        try:
            parent.remove(tbl)
            p41._p.addnext(tbl)
        except Exception:
            pass

    def map_header_columns(self, row):
        mapping = {}
        for i, cell in enumerate(row.cells):
            h = normalize_text(cell.text)
            if "LOTE" in h:
                mapping["lote"] = i
            elif "SIPLAN" in h or ("COD" in h and "siplan" not in mapping):
                mapping["siplan"] = i
            elif h == "ITEM":
                mapping["item"] = i
            elif "QUANT" in h:
                mapping["quant"] = i
            elif "DESCRI" in h:
                mapping["descricao"] = i
            elif "APRESENT" in h:
                mapping["apresentacao"] = i
            elif "MARCA" in h:
                mapping["marca"] = i
            elif "MODELO" in h:
                mapping["modelo"] = i
            elif "UNIT" in h or ("PRECO" in h and "TOTAL" not in h):
                mapping["preco_unitario"] = i
            elif "TOTAL" in h or "PRECO TOTAL" in h:
                mapping["preco_total"] = i
        return mapping

    def put(self, row, mapping, key, value, description=False):
        if key not in mapping:
            return
        idx = mapping[key]
        if idx >= len(row.cells):
            return
        if description:
            self.set_description_cell(row.cells[idx], value)
        else:
            self.set_cell_text_keep_style(row.cells[idx], value)

    def find_price_table(self, doc):
        for t in doc.tables:
            if not t.rows:
                continue
            h = normalize_text(" ".join(c.text for c in t.rows[0].cells))
            if (("SIPLAN" in h or "LOTE" in h) and "DESCRI" in h and ("PRECO" in h or "VALOR" in h)):
                return t
        return None

    def write_total(self, row, mapping, total):
        cells = row.cells
        if not cells:
            return

        last_idx = mapping.get("preco_total", len(cells) - 1)
        last_idx = min(last_idx, len(cells) - 1)

        try:
            label = cells[0]
            for i in range(1, last_idx):
                label = label.merge(cells[i])
            self.set_cell_text_keep_style(label, "VALOR TOTAL:")
            self.set_cell_text_keep_style(cells[last_idx], format_money(total, 2))
            self.center(label)
            self.center(cells[last_idx])
            self.bold(label)
            self.bold(cells[last_idx])
            self.horizontal(label)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Cm(0.55)
        except Exception:
            self.set_cell_text_keep_style(cells[0], "VALOR TOTAL:")
            self.set_cell_text_keep_style(cells[-1], format_money(total, 2))

    def replace_total_paragraph(self, doc, ata):
        txt = f"4.2 Valor total dos preços registrados: {format_money(ata.valor_total, 2)} ({valor_extenso(ata.valor_total)})."
        for p in doc.paragraphs:
            if "Valor total dos preços registrados" in p.text:
                self.set_paragraph_text_keep_style(p, txt)
                return

    # ------------------- APÊNDICE -------------------
    def replace_appendix(self, doc, ata):
        table = self.find_appendix_table(doc)
        if table is None:
            return

        data_idx = 1 if len(table.rows) > 1 else 0
        data_template = deepcopy(table.rows[data_idx]._tr)

        while len(table.rows) > 1:
            table.rows[-1]._tr.getparent().remove(table.rows[-1]._tr)

        for item in ata.itens:
            table._tbl.append(deepcopy(data_template))
            row = table.rows[-1]
            self.clear_row_text_only(row)

            values = item.appendix_cells_text or [
                item.codigo_siplan,
                str(item.numero_item),
                item.descricao_oficial,
                item.apresentacao,
                format_qty(item.quantidade),
            ]

            for i, v in enumerate(values):
                if i < len(row.cells):
                    if i == 2:
                        self.set_description_cell(row.cells[i], v)
                    else:
                        self.set_cell_text_keep_style(row.cells[i], v)

    def find_appendix_table(self, doc):
        best = None
        score_best = 0
        for t in doc.tables:
            if not t.rows:
                continue
            h = normalize_text(" ".join(c.text for c in t.rows[0].cells))
            score = sum(1 for w in ["SIPLAN", "LOTE", "ITEM", "DESCRITIVO", "APRESENT", "ICISMEP", "TOTAL"] if w in h)
            if len(t.rows[0].cells) > 10:
                score += 2
            if score > score_best:
                best, score_best = t, score
        return best if score_best >= 5 else None

    def replace_signature(self, doc, ata):
        txt = f"{self.val(ata.representante)}\n{self.clean_supplier_name(ata.fornecedor_nome).title()}"
        if len(doc.tables) >= 3:
            try:
                self.set_cell_text_keep_style(doc.tables[2].rows[0].cells[-1], txt)
            except Exception:
                pass

    # ------------------- PRESERVAÇÃO DE ESTILO/COR -------------------
    def clear_row_text_only(self, row):
        for c in row.cells:
            self.set_cell_text_keep_style(c, "")

    def set_cell_text_keep_style(self, cell, text):
        text = "" if text is None else str(text)

        if cell.paragraphs:
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].text = text
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(text)

            for pp in cell.paragraphs[1:]:
                for r in pp.runs:
                    r.text = ""
        else:
            cell.text = text

        for p in cell.paragraphs:
            for r in p.runs:
                if r.font.size is None:
                    r.font.size = Pt(8)
                if r.font.name is None:
                    r.font.name = "Arial"

    def set_description_cell(self, cell, text):
        text = "" if text is None else str(text)
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

        for r in list(p.runs):
            r.text = ""
        for extra_p in cell.paragraphs[1:]:
            for r in extra_p.runs:
                r.text = ""

        if " - " in text:
            first, rest = text.split(" - ", 1)
            r1 = p.add_run(first)
            r1.bold = True
            r2 = p.add_run(" - " + rest)
            r2.bold = False
            runs = [r1, r2]
        else:
            r = p.add_run(text)
            r.bold = True
            runs = [r]

        for r in runs:
            r.font.name = "Arial"
            r.font.size = Pt(8)

    def set_paragraph_text_keep_style(self, p, text):
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(text)

    def center(self, cell):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def bold(self, cell):
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    def horizontal(self, cell):
        tcPr = cell._tc.get_or_add_tcPr()
        for ch in list(tcPr):
            if ch.tag == qn("w:textDirection"):
                tcPr.remove(ch)
        el = OxmlElement("w:textDirection")
        el.set(qn("w:val"), "lrTb")
        tcPr.append(el)
