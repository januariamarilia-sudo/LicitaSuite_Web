from pathlib import Path
import shutil
import re


HELPER_CODE = r'''
def formatar_quantidade_br(valor):
    if valor is None:
        return ""

    texto = str(valor).strip()

    if not texto:
        return ""

    texto_limpo = texto.replace(".", "").replace(",", "").strip()

    try:
        numero = int(float(texto_limpo))
        return f"{numero:,}".replace(",", ".")
    except Exception:
        return texto


def aplicar_formatacao_assinatura(documento):
    try:
        from docx.shared import Pt
    except Exception:
        return documento

    palavras_chave = [
        "januária", "januaria", "medeiros",
        "icismep", "consórcio", "consorcio",
        "gerenciador", "órgão gerenciador", "orgao gerenciador",
    ]

    def aplicar_em_paragrafo(paragrafo):
        texto = (paragrafo.text or "").strip()
        texto_norm = texto.lower()

        if not texto:
            return

        eh_linha_assinatura = (
            len(texto) <= 120
            and any(p in texto_norm for p in palavras_chave)
            and not texto_norm.startswith("cláusula")
            and not texto_norm.startswith("4.")
            and not texto_norm.startswith("dos preços")
        )

        if not eh_linha_assinatura:
            return

        for run in paragrafo.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.bold = True

            if run.text and run.text.isupper() and len(run.text.split()) > 1:
                run.text = run.text.title()

    for paragrafo in documento.paragraphs:
        aplicar_em_paragrafo(paragrafo)

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    aplicar_em_paragrafo(paragrafo)

    return documento
'''


def localizar_candidatos(root):
    candidatos = []

    for path in root.rglob("*.py"):
        if "__pycache__" in path.parts:
            continue

        txt = path.read_text(encoding="utf-8", errors="ignore")
        nome = str(path).lower()

        score = 0
        if "docx" in txt.lower() or "Document(" in txt:
            score += 2
        if "quantidade" in txt.lower() or "quantidade_pdf" in txt.lower():
            score += 2
        if "valor_unitario" in txt.lower() or "valor_total" in txt.lower():
            score += 2
        if "copy_model_generator" in nome or "generator" in nome or "docx_engine" in nome:
            score += 3

        if score >= 5:
            candidatos.append((score, path))

    candidatos.sort(reverse=True, key=lambda x: x[0])
    return candidatos


def inserir_helpers(text):
    if "def formatar_quantidade_br(" in text:
        return text, False

    lines = text.splitlines()
    idx = 0

    for i, line in enumerate(lines):
        if line.startswith("import ") or line.startswith("from "):
            idx = i + 1

    lines.insert(idx, HELPER_CODE)
    return "\n".join(lines) + "\n", True


def trocar_quantidades(text):
    changed = False

    replacements = [
        ("str(item.quantidade_pdf)", "formatar_quantidade_br(item.quantidade_pdf)"),
        ("str(item.quantidade)", "formatar_quantidade_br(item.quantidade)"),
        ("str(item.get('quantidade'))", "formatar_quantidade_br(item.get('quantidade'))"),
        ('str(item.get("quantidade"))', 'formatar_quantidade_br(item.get("quantidade"))'),
        ("str(item.get('quantidade_pdf'))", "formatar_quantidade_br(item.get('quantidade_pdf'))"),
        ('str(item.get("quantidade_pdf"))', 'formatar_quantidade_br(item.get("quantidade_pdf"))'),
        ("str(quantidade)", "formatar_quantidade_br(quantidade)"),
        ("str(qtd)", "formatar_quantidade_br(qtd)"),
    ]

    for old, new in replacements:
        if old in text and new not in text:
            text = text.replace(old, new)
            changed = True

    return text, changed


def inserir_assinatura_antes_do_save(text):
    if "aplicar_formatacao_assinatura(document)" in text or "aplicar_formatacao_assinatura(doc)" in text or "aplicar_formatacao_assinatura(documento)" in text:
        return text, False

    for var_name in ["document", "doc", "documento"]:
        pattern = f"{var_name}.save("
        idx = text.find(pattern)
        if idx != -1:
            line_start = text.rfind("\n", 0, idx) + 1
            indent = text[line_start:idx]
            call = f"{indent}{var_name} = aplicar_formatacao_assinatura({var_name})\n"
            text = text[:line_start] + call + text[line_start:]
            return text, True

    return text, False


def patch_file(path):
    text = path.read_text(encoding="utf-8", errors="ignore")
    original = text

    text, c1 = inserir_helpers(text)
    text, c2 = trocar_quantidades(text)
    text, c3 = inserir_assinatura_antes_do_save(text)

    if text != original and (c1 or c2 or c3):
        backup = path.with_suffix(path.suffix + ".bak_3_9_formatacao")
        if not backup.exists():
            shutil.copy2(path, backup)
        path.write_text(text, encoding="utf-8")
        return True

    return False


def main():
    root = Path.cwd()
    candidatos = localizar_candidatos(root)

    if not candidatos:
        print("ERRO: não localizei automaticamente o gerador DOCX.")
        print("Envie o resultado de:")
        print('grep -R "quantidade_pdf\\|valor_unitario\\|valor_total" licitasuite -n')
        return

    aplicados = []

    for score, path in candidatos[:6]:
        if patch_file(path):
            aplicados.append(path)

    if not aplicados:
        print("ATENÇÃO: encontrei candidatos, mas não consegui aplicar alteração automática.")
        print("Candidatos encontrados:")
        for score, path in candidatos[:8]:
            print(f"- {path} | score {score}")
        print("")
        print("Envie o arquivo principal do gerador DOCX para ajuste manual.")
        return

    print("PATCH APLICADO COM SUCESSO.")
    print("Arquivos alterados:")
    for p in aplicados:
        print(f"- {p}")

    print("")
    print("Agora rode:")
    print("git add licitasuite docs/LICITASUITE_WEB_3_9_FORMATACAO_QUANTIDADE_ASSINATURA.md aplicar_patch_formatacao_3_9.py")
    print('git commit -m "Formata quantidade e assinatura das atas"')
    print("git push")


if __name__ == "__main__":
    main()
